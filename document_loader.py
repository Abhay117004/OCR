from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import cv2
import docx
import numpy as np
import pdfplumber
import pypdfium2 as pdfium


@dataclass
class NativeDocument:
    text: str
    pages: list[dict]
    metadata: dict


@dataclass
class RasterDocument:
    pages: list[np.ndarray]
    metadata: dict


def load_document(path: str | Path) -> NativeDocument | RasterDocument:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        image = cv2.imread(str(path))
        if image is None:
            raise FileNotFoundError(f"Cannot read image: {path}")
        return RasterDocument(
            pages=[cv2.cvtColor(image, cv2.COLOR_BGR2RGB)],
            metadata={"source_type": "image", "path": str(path)},
        )

    if suffix == ".pdf":
        return _load_pdf(path)

    if suffix == ".docx":
        return _load_docx(path)

    if suffix == ".doc":
        return _load_doc(path)

    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8")
        return NativeDocument(
            text=text,
            pages=[{"page": 1, "blocks": [{"type": "paragraph", "text": text}]}],
            metadata={"source_type": "text", "path": str(path)},
        )

    raise ValueError(
        f"Unsupported file type: {suffix}. Supported: images, PDF, DOCX, TXT, MD."
    )


def _load_pdf(path: Path) -> NativeDocument | RasterDocument:
    pages: list[dict] = []
    native_chunks: list[str] = []

    with pdfplumber.open(str(path)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            words = page.extract_words() or []
            pages.append({"page": page_index, "text": text, "words": words})
            if text:
                native_chunks.append(text)

    if native_chunks:
        return NativeDocument(
            text="\n\n".join(native_chunks),
            pages=pages,
            metadata={"source_type": "pdf-native", "path": str(path)},
        )

    pdf_doc = pdfium.PdfDocument(str(path))
    images: list[np.ndarray] = []
    for page_index in range(len(pdf_doc)):
        bitmap = pdf_doc[page_index].render(scale=2.0).to_pil()
        images.append(np.array(bitmap.convert("RGB")))

    return RasterDocument(
        pages=images,
        metadata={"source_type": "pdf-raster", "path": str(path)},
    )


def _load_docx(path: Path) -> NativeDocument:
    document = docx.Document(str(path))
    blocks: list[dict] = []
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        blocks.append({"type": "paragraph", "text": text})
        text_parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if not cells:
                continue
            text = " | ".join(cells)
            blocks.append({"type": "table_row", "text": text, "cells": cells})
            text_parts.append(text)

    return NativeDocument(
        text="\n".join(text_parts),
        pages=[{"page": 1, "blocks": blocks}],
        metadata={"source_type": "docx", "path": str(path)},
    )


def _load_doc(path: Path) -> NativeDocument:
    try:
        import tempfile
        from win32com.client import Dispatch
    except Exception as exc:
        raise ValueError(
            ".doc support requires Microsoft Word on Windows plus pywin32."
        ) from exc

    temp_dir = Path(tempfile.mkdtemp(prefix="ocr_doc_"))
    docx_path = temp_dir / f"{path.stem}.docx"
    word = None
    document = None
    try:
        word = Dispatch("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(path.resolve()))
        document.SaveAs(str(docx_path), FileFormat=16)
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()

    return _load_docx(docx_path)
